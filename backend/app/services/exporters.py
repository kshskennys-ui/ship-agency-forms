import json
import html
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from ..paths import EXPORT_DIR, HELPER_DIR, TEMPLATE_DIR
from .forecast import format_crew_change


def _helper_path(name):
    packaged = HELPER_DIR / name
    return packaged if packaged.exists() else Path(__file__).with_name(name)


def _text(value):
    return "" if value is None else str(value)


def _cn_time(value):
    return f"{value.month}月{value.day}日{value:%H%M}时" if value else ""


def _date_text(value, separator="/"):
    return value.strftime(f"%Y{separator}%m{separator}%d") if hasattr(value, "strftime") else _text(value)


def _flight_time(value):
    return value.strftime("%m/%d %H：%M") if value else ""


def _berth_phase(value):
    raw = _text(value).strip()
    match = re.search(r"([一二三四五六七八九十百]+期|\d+期)", raw)
    return match.group(1) if match else raw


def export_inbound_form(vessel, voyage, crew, form_type):
    """Export the selected inbound Excel form from the supplied workbook template."""
    extra = json.loads(voyage.extra_json or "{}")
    vessel_extra = json.loads(vessel.extra_json or "{}")
    nationality_counts = Counter(person.nationality or "待人工填写" for person in crew)
    nationality_distribution = "\n".join(
        f"{nationality}{count}名" for nationality, count in nationality_counts.items()
    )
    payload = {
        "vessel": {
            "imo": vessel.imo,
            "chinese_name": vessel.chinese_name,
            "english_name": vessel.english_name,
            "nationality": vessel.nationality,
            "call_sign": vessel.call_sign,
            "shipping_company": vessel.shipping_company,
            "mmsi": vessel.mmsi,
            "gross_tonnage": vessel.gross_tonnage,
            "net_tonnage": vessel.net_tonnage,
            "nationality_certificate_no": (
                vessel_extra.get("nationality_certificate_no")
                or vessel_extra.get("nationality_certificate_number")
                or vessel_extra.get("registry_certificate_no")
                or ""
            ),
        },
        "voyage": {
            "inbound_voyage_no": voyage.inbound_voyage_no,
            "outbound_voyage_no": voyage.outbound_voyage_no,
            "arrival_time": voyage.arrival_time.isoformat() if voyage.arrival_time else None,
            "berth": _berth_phase(voyage.berth),
            "previous_port": voyage.previous_port,
            "next_port": voyage.next_port,
            "next_port_country": voyage.next_port_country,
            "extra": extra,
        },
        "crew": {
            "captain": crew[0].name if crew else "",
            "total_count": len(crew),
            "passenger_count": extra.get("passenger_count", 0),
            "nationality_distribution": nationality_distribution,
        },
    }
    bundled_node = Path(r"C:\Users\UA\.cache\codex-runtimes\codex-primary-runtime\dependencies\node\bin\node.exe")
    node_path = os.getenv("SHIP_AGENCY_NODE") or (str(bundled_node) if bundled_node.exists() else None) or shutil.which("node")
    helper = _helper_path("inbound_forms_exporter.mjs")
    template = TEMPLATE_DIR / "inbound_forms_template.xlsx"
    label = "强文献总申" if form_type == "general" else "海关货申"
    output = EXPORT_DIR / f"{label}_{voyage.id}_{datetime.now():%Y%m%d%H%M%S}.xlsx"
    with tempfile.NamedTemporaryFile("w", encoding="utf-8", suffix=".json", delete=False) as payload_file:
        json.dump(payload, payload_file, ensure_ascii=False)
        payload_path = Path(payload_file.name)
    try:
        subprocess.run(
            [node_path, str(helper), str(template), str(output), str(payload_path), form_type],
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=90,
        )
    except subprocess.CalledProcessError as exc:
        message = (exc.stderr or exc.stdout or "未知错误").strip()
        raise RuntimeError(f"{label}生成失败：{message}") from exc
    finally:
        payload_path.unlink(missing_ok=True)
    return output


def _resize_crew_table(table, people_count, header_rows=2):
    """Keep the fixed title/header rows and exactly one data row per crew member."""
    desired_rows = people_count + header_rows
    while len(table.rows) > desired_rows:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < desired_rows:
        table.add_row()


def _set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _set_paragraph_text_preserve_images(paragraph, text):
    """Replace visible text without deleting legacy VML images in the paragraph."""
    image_runs = [run for run in paragraph.runs if "w:pict" in run._r.xml or "w:drawing" in run._r.xml]
    if not image_runs:
        _set_paragraph_text(paragraph, text)
        return
    text_runs = [run for run in paragraph.runs if run not in image_runs]
    target = next((run for run in reversed(text_runs) if run.text.strip()), None)
    if target is None:
        target = text_runs[-1] if text_runs else paragraph.add_run()
    target.text = text


def _set_customs_agent_identity(document, company, company_en):
    """Replace the sample company identity in the customs template header/body."""
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "示例船务代理有限公司":
            _set_paragraph_text_preserve_images(paragraph, company)
        elif text.lower() == "example shipping agency co., ltd.":
            _set_paragraph_text(paragraph, company_en)
    for section in document.sections:
        for paragraph in (*section.header.paragraphs, *section.footer.paragraphs):
            text = paragraph.text.strip()
            if text == "示例船务代理有限公司":
                _set_paragraph_text_preserve_images(paragraph, company)
            elif text.lower() == "example shipping agency co., ltd.":
                _set_paragraph_text(paragraph, company_en)


def _set_customs_header(document, company, company_en, address_zh, address_en, phone, fax, email):
    """Apply the actual agency letterhead from the approved customs example."""
    header_values = {
        # Paragraph 0 already contains the leading spacing in its separate
        # text run next to the Logo; keep that run intact.
        0: company,
        1: f"      {company_en} ",
        3: f"{address_zh} ",
        4: address_en,
        5: f"电话(Tel):86-020-{phone}    传真(Fax):86-020-{fax}   E-mail:{email} ",
    }
    for index, value in header_values.items():
        if index < len(document.paragraphs):
            _set_paragraph_text_preserve_images(document.paragraphs[index], value)


def _apply_table_borders(table):
    """Ensure every cell in a generated crew table has the visible template grid."""
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = qn(f"w:{edge}")
                element = borders.find(tag)
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "4")
                element.set(qn("w:space"), "0")
                element.set(qn("w:color"), "000000")


def _crew_document_type(person):
    extra = {}
    try:
        extra = json.loads(getattr(person, "extra_json", "{}") or "{}")
    except (TypeError, json.JSONDecodeError):
        extra = {}
    explicit = extra.get("document_type") or extra.get("document_category")
    raw_values = extra.get("raw_row") or []
    raw_text = " ".join(_text(value) for value in raw_values).lower()
    source_text = f"{_text(explicit)} {raw_text}"
    if "海员证" in source_text or "seaman" in source_text or "seafarer" in source_text:
        return "海员证"
    if "护照" in source_text or "passport" in source_text:
        return "护照"
    return _text(explicit or "其他证件").strip() or "其他证件"


def _crew_document_summary(crew):
    counts = Counter(_crew_document_type(person) for person in crew)
    ordered_types = ["海员证", "护照"] + sorted(
        document_type for document_type in counts if document_type not in {"海员证", "护照"}
    )
    return "".join(
        f"{counts[document_type]}{document_type}"
        for document_type in ordered_types
        if counts.get(document_type)
    )


def _patch_docx_paragraph_text(document_xml, marker, replacement):
    paragraph_pattern = re.compile(r"(<w:p\b[^>]*>)(.*?)(</w:p>)", re.S)
    text_pattern = re.compile(r"(<w:t\b[^>]*>)(.*?)(</w:t>)", re.S)
    replaced = False

    def paragraph_replacer(match):
        nonlocal replaced
        opening, body, closing = match.groups()
        visible_text = "".join(
            html.unescape(text_match.group(2))
            for text_match in text_pattern.finditer(body)
        )
        if marker not in visible_text:
            return match.group(0)

        first_text = True

        def text_replacer(text_match):
            nonlocal first_text
            if first_text:
                first_text = False
                return (
                    f"{text_match.group(1)}"
                    f"{xml_escape(replacement)}"
                    f"{text_match.group(3)}"
                )
            return f"{text_match.group(1)}{text_match.group(3)}"

        replaced = True
        return f"{opening}{text_pattern.sub(text_replacer, body)}{closing}"

    patched = paragraph_pattern.sub(paragraph_replacer, document_xml)
    if not replaced:
        raise ValueError(f"未找到外勤收据模板字段：{marker}")
    return patched


def export_outer_field_receipt(vessel, voyage, crew):
    """Export the outer-field received list using the supplied Word template."""
    ship_name = "/".join(
        name for name in [_text(vessel.chinese_name), _text(vessel.english_name)] if name
    )
    template = TEMPLATE_DIR / "outer_field_receipt_template.docx"
    output = EXPORT_DIR / f"外勤收据_{voyage.id}_{datetime.now():%Y%m%d%H%M%S}.docx"
    with zipfile.ZipFile(template, "r") as source_zip, zipfile.ZipFile(
        output, "w", compression=zipfile.ZIP_DEFLATED
    ) as output_zip:
        for entry in source_zip.infolist():
            content = source_zip.read(entry.filename)
            if entry.filename == "word/document.xml":
                document_xml = content.decode("utf-8")
                document_xml = _patch_docx_paragraph_text(
                    document_xml, "船名：", f"船名：{ship_name}"
                )
                document_xml = _patch_docx_paragraph_text(
                    document_xml,
                    "PASSPORT",
                    f"证件类别：{_crew_document_summary(crew)}",
                )
                content = document_xml.encode("utf-8")
            output_zip.writestr(entry, content)
    return output


def _is_chinese_nationality(value):
    normalized = _text(value).strip().lower()
    return normalized in {"中国", "中华人民共和国", "china", "cn"} or normalized.startswith("中国")


def _is_hk_macao_taiwan_nationality(value):
    normalized = _text(value).strip().lower()
    return any(token in normalized for token in ("香港", "澳门", "台湾", "hong kong", "macao", "taiwan"))


def _border_document_counts(crew):
    counts = {
        "chinese_seaman": 0,
        "foreign_passport": 0,
        "foreign_seaman": 0,
        "hk_macao_taiwan": 0,
    }
    for person in crew:
        if _is_hk_macao_taiwan_nationality(person.nationality):
            counts["hk_macao_taiwan"] += 1
            continue
        document_type = _crew_document_type(person)
        if document_type == "海员证":
            key = "chinese_seaman" if _is_chinese_nationality(person.nationality) else "foreign_seaman"
        elif document_type == "护照":
            key = "foreign_passport"
        else:
            key = "chinese_seaman" if _is_chinese_nationality(person.nationality) else "foreign_passport"
        counts[key] += 1
    return counts


def _border_datetime(value):
    if not value:
        return ""
    return f"{value.year}年{value.month}月{value.day}日 {value:%H:%M}"


def export_border_inspection(vessel, voyage, crew, changes):
    """Export the border-inspection procedures sheet while preserving its source layout."""
    bundled_node = Path(r"C:\Users\UA\.cache\codex-runtimes\codex-primary-runtime\dependencies\node\bin\node.exe")
    node_path = os.getenv("SHIP_AGENCY_NODE") or (str(bundled_node) if bundled_node.exists() else None) or shutil.which("node")
    helper = _helper_path("border_inspection_exporter.mjs")
    template = TEMPLATE_DIR / "border_inspection_procedures_template.xlsx"
    output = EXPORT_DIR / f"边检手续表_{voyage.id}_{datetime.now():%Y%m%d%H%M%S}.xlsx"
    document_counts = _border_document_counts(crew)
    female_count = sum(1 for person in crew if _text(person.gender).strip() in {"女", "female", "f"})
    current_change = bool(changes)
    change_summary = format_crew_change(changes) if current_change else ""
    ports = [value for value in [voyage.previous_port, "南沙", voyage.next_port] if value]
    payload = {
        "vessel": {
            "chinese_name": vessel.chinese_name,
            "english_name": vessel.english_name,
            "shipping_company": vessel.shipping_company,
        },
        "voyage": {
            "arrival_time": _border_datetime(voyage.arrival_time),
            "departure_time": _border_datetime(voyage.departure_time),
            "previous_port": voyage.previous_port,
            "next_port": voyage.next_port,
            "port_sequence": "-".join(ports) + ("；" if ports else ""),
            "berth": voyage.berth,
        },
        "crew": {
            **document_counts,
            "female_count": female_count,
        },
        "changes": {
            "has_current": current_change,
            "current_summary": change_summary,
            "has_domestic": False,
            "domestic_summary": "",
            "other_summary": "",
        },
    }
    with tempfile.NamedTemporaryFile("w", encoding="utf-8", suffix=".json", delete=False) as payload_file:
        json.dump(payload, payload_file, ensure_ascii=False)
        payload_path = Path(payload_file.name)
    try:
        subprocess.run(
            [node_path, str(helper), str(template), str(output), str(payload_path)],
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=90,
        )
    except subprocess.CalledProcessError as exc:
        message = (exc.stderr or exc.stdout or "未知错误").strip()
        raise RuntimeError(f"边检手续表生成失败：{message}") from exc
    finally:
        payload_path.unlink(missing_ok=True)
    return output


def export_tonnage(vessel, voyage, application):
    bundled_node = Path(r"C:\Users\UA\.cache\codex-runtimes\codex-primary-runtime\dependencies\node\bin\node.exe")
    node_path = os.getenv("SHIP_AGENCY_NODE") or (str(bundled_node) if bundled_node.exists() else None) or shutil.which("node")
    helper = _helper_path("tonnage_exporter.mjs")
    template = TEMPLATE_DIR / "tonnage_template.xlsx"
    duration_label = {30: "三十天期（30days）", 90: "九十天期（90days）", 365: "一年期（1year）"}.get(application.duration_days, "待选择")
    output = EXPORT_DIR / f"tonnage_{voyage.id}_{datetime.now():%Y%m%d%H%M%S}.xlsx"
    today = datetime.now().date()
    payload = {
        "vessel": {
            "imo": vessel.imo,
            "chinese_name": vessel.chinese_name,
            "english_name": vessel.english_name,
            "nationality": vessel.nationality,
            "net_tonnage": vessel.net_tonnage,
        },
        "voyage": {
            "arrival_date": voyage.arrival_time.date().isoformat() if voyage.arrival_time else None,
        },
        "application": {
            "amount": application.amount,
            "pre_entry_no": application.pre_entry_no,
            "charter_relation": application.charter_relation or "其他",
            "duration_label": duration_label,
        },
        "today": {
            "year": today.year,
            "month": today.month,
            "day": today.day,
        },
    }
    with tempfile.NamedTemporaryFile("w", encoding="utf-8", suffix=".json", delete=False) as payload_file:
        json.dump(payload, payload_file, ensure_ascii=False)
        payload_path = Path(payload_file.name)
    try:
        subprocess.run(
            [node_path, str(helper), str(template), str(output), str(payload_path)],
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=90,
        )
    except subprocess.CalledProcessError as exc:
        message = (exc.stderr or exc.stdout or "未知错误").strip()
        raise RuntimeError(f"吨税表生成失败：{message}") from exc
    finally:
        payload_path.unlink(missing_ok=True)
    return output


def export_health_declaration(vessel, voyage, crew):
    extra = json.loads(voyage.extra_json or "{}")
    ship_name = "/".join(x for x in [vessel.chinese_name, vessel.english_name] if x)
    voyage_number = "/".join(x for x in [voyage.inbound_voyage_no, voyage.outbound_voyage_no] if x)
    port_sequence = extra.get("port_sequence") or "-".join(
        x for x in [voyage.previous_port, "南沙", voyage.next_port] if x
    )
    crew_rows = []
    for person in crew:
        person_extra = json.loads(getattr(person, "extra_json", "{}") or "{}")
        crew_rows.append(
            {
                "name": person.name,
                "document_no": person.document_no,
                "gender": person.gender,
                "birth_date": person.birth_date.isoformat() if person.birth_date else None,
                "temperature": person_extra.get("temperature", ""),
            }
        )
    declaration_date = voyage.arrival_time.date().isoformat() if voyage.arrival_time else datetime.now().date().isoformat()
    payload = {
        "ship_name": ship_name,
        "imo": vessel.imo or "",
        "voyage_number": voyage_number,
        "declaration_date": declaration_date,
        "recent_places": port_sequence,
        "contact_default": "No",
        "symptom_default": "No",
        "crew": crew_rows,
    }
    bundled_node = Path(r"C:\Users\UA\.cache\codex-runtimes\codex-primary-runtime\dependencies\node\bin\node.exe")
    node_path = os.getenv("SHIP_AGENCY_NODE") or (str(bundled_node) if bundled_node.exists() else None) or shutil.which("node")
    helper = _helper_path("health_declaration_exporter.mjs")
    template = TEMPLATE_DIR / "health_declaration_template.xlsx"
    output = EXPORT_DIR / f"health_declaration_{voyage.id}_{datetime.now():%Y%m%d%H%M%S}.xlsx"
    with tempfile.NamedTemporaryFile("w", encoding="utf-8", suffix=".json", delete=False) as payload_file:
        json.dump(payload, payload_file, ensure_ascii=False)
        payload_path = Path(payload_file.name)
    try:
        subprocess.run(
            [node_path, str(helper), str(template), str(output), str(payload_path)],
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=60,
        )
    except subprocess.CalledProcessError as exc:
        message = (exc.stderr or exc.stdout or "未知错误").strip()
        raise RuntimeError(f"体温和健康申报表生成失败：{message}") from exc
    finally:
        payload_path.unlink(missing_ok=True)
    return output


def export_crew_change(vessel, voyage, people):
    document = Document(TEMPLATE_DIR / "crew_change_template.docx")
    extra = json.loads(voyage.extra_json or "{}")
    agent_company = extra.get("agent_company", "广州港中联国际船务代理有限公司")
    agent_name = extra.get("agent_name", "")
    agent_contact = extra.get("agent_contact", "")
    ship_type = extra.get("ship_type", "国际航行船舶")
    port_sequence = extra.get("port_sequence") or "-".join(x for x in [voyage.previous_port, "南沙", voyage.next_port] if x)
    ship_label = "/".join(x for x in [vessel.english_name, vessel.chinese_name] if x)
    if document.paragraphs:
        title_paragraph = document.paragraphs[0]
        if len(title_paragraph.runs) >= 7:
            title_paragraph.runs[0].text = "“"
            title_paragraph.runs[1].text = vessel.english_name or ""
            title_paragraph.runs[2].text = "/" if vessel.english_name and vessel.chinese_name else ""
            title_paragraph.runs[3].text = vessel.chinese_name or ""
            title_paragraph.runs[4].text = "”"
            title_paragraph.runs[5].text = "船员换班申请"
            title_paragraph.runs[6].text = "表"
        else:
            _set_paragraph_text(title_paragraph, f"“{ship_label}”船员换班申请表")
    company_paragraph = next((p for p in document.paragraphs if "公司" in p.text and "船员" not in p.text), None)
    if company_paragraph:
        _set_paragraph_text(company_paragraph, agent_company)
        if company_paragraph.runs:
            company_paragraph.runs[0].font.size = Pt(10)
    date_paragraph = next((p for p in document.paragraphs if "." in p.text and p is not document.paragraphs[0]), None)
    if date_paragraph:
        _set_paragraph_text(date_paragraph, datetime.now().strftime("%Y.%m.%d"))
    info = document.tables[0]
    rows = [
        [ship_label, vessel.imo or "", vessel.nationality or "", agent_company, f"{agent_name} {agent_contact}".strip()],
        [voyage.berth or "", _cn_time(voyage.arrival_time), _cn_time(voyage.departure_time), ship_type, port_sequence],
    ]
    for row_index, values in zip((1, 3), rows):
        for cell, value in zip(info.rows[row_index].cells, values):
            cell.text = _text(value)
    for cell in info.rows[3].cells[1:3]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
    down = [p for p in people if p.direction == "down"]
    up = [p for p in people if p.direction == "up"]
    total_change = len(up) + len(down)
    narrative = f"因船员有离船公休和登轮接班需求，现为{total_change}名船员申请办理船员离船公休和登轮接班手续（{len(up)}上{len(down)}下），{_cn_time(voyage.arrival_time)}过船员通道，我司承诺船员健康情况正常，船员上下轮均不携带违法违规物品。"
    for cell in info.rows[4].cells:
        cell.text = narrative
    up_table = document.tables[1]
    _resize_crew_table(up_table, len(up))
    for index, person in enumerate(up, start=2):
        values = [index - 1, person.name, person.nationality, person.gender, _date_text(person.birth_date), person.document_no, ""]
        for cell, value in zip(up_table.rows[index].cells, values):
            cell.text = _text(value)

    table = document.tables[2]
    _resize_crew_table(table, len(down))
    for index, person in enumerate(down, start=2):
        values = [index - 1, person.name, person.nationality, person.gender, _date_text(person.birth_date, "-"), person.document_no, person.reason or "", "是" if person.temporary_entry_permit else "否", person.flight_no or "", _flight_time(person.flight_time), person.route or ""]
        for cell, value in zip(table.rows[index].cells, values):
            cell.text = _text(value)
    output = EXPORT_DIR / f"crew_change_{voyage.id}_{datetime.now():%Y%m%d%H%M%S}.docx"
    document.save(output)
    return output


def export_crew_change_customs(vessel, voyage, people):
    """Export the customs-style crew change application from its supplied Word template."""
    document = Document(TEMPLATE_DIR / "customs_crew_change_template.docx")
    extra = json.loads(voyage.extra_json or "{}")
    agent_company = extra.get("agent_company", "广州港中联国际船务代理有限公司")
    agent_company_en = extra.get(
        "agent_company_en",
        "Guangzhou Port Unitrans Agency Co., Ltd.",
    )
    agent_address_zh = extra.get(
        "agent_address_zh",
        "中国 广州 南沙区万顷沙龙穴岛 口岸大厦17楼  邮编:511462",
    )
    agent_address_en = extra.get(
        "agent_address_en",
        "Floor17,Kouan Building,Longxue Islan,Nansha District,Guangzhou,P.R.C.",
    )
    agent_phone = extra.get("agent_phone", "39080621")
    agent_fax = extra.get("agent_fax", "34660550")
    agent_email = extra.get("agent_email", "nsa.shipping@unitrans-agency.com")
    agent_name = extra.get("agent_name", "")
    agent_contact = extra.get("agent_contact", "")
    ship_type = extra.get("ship_type", "集装箱货船")
    port_sequence = extra.get("port_sequence") or "-".join(
        x for x in [voyage.previous_port, "南沙", voyage.next_port] if x
    )
    chinese_name = _text(vessel.chinese_name)
    english_name = _text(vessel.english_name)
    up = [p for p in people if p.direction == "up"]
    down = [p for p in people if p.direction == "down"]
    _set_customs_agent_identity(document, agent_company, agent_company_en)
    _set_customs_header(
        document,
        agent_company,
        agent_company_en,
        agent_address_zh,
        agent_address_en,
        agent_phone,
        agent_fax,
        agent_email,
    )

    title = next((p for p in document.paragraphs if "换班申请" in p.text), None)
    if title:
        _set_paragraph_text(title, f"“{chinese_name}”南沙港换班申请")

    body = next((p for p in document.paragraphs if p.text.startswith("兹有我司代理")), None)
    if body:
        arrival = _customs_datetime(voyage.arrival_time, include_year=True)
        departure = _customs_datetime(voyage.departure_time, include_year=False)
        body_text = (
            f"兹有我司代理{_text(vessel.nationality)}籍{ship_type}“{chinese_name}”，"
            f"英文名：{english_name},IMO：{_text(vessel.imo)}，预计于{arrival}靠泊"
            f"{_text(voyage.berth or '南沙港')}，计划于{departure}离泊。该轮挂靠港序：{port_sequence}"
        )
        _set_paragraph_text(body, body_text)

    change_request = next((p for p in document.paragraphs if p.text.startswith("受船东委托")), None)
    if change_request:
        _set_paragraph_text(
            change_request,
            f"受船东委托，根据船员管理安排，需在本港安排船员更动{len(up)}上{len(down)}下，船员信息如下：",
        )

    closing = next((p for p in document.paragraphs if p.text.startswith("就以上信息")), None)
    if closing:
        application_time = extra.get("crew_change_time") or voyage.arrival_time
        _set_paragraph_text(
            closing,
            f"就以上信息，特向贵部门申请于{_customs_datetime(application_time, include_year=False)}时在船员通道办理船员更动{len(up)}上{len(down)}下手续， 感谢贵部门对我司的大力支持！",
        )

    sign_company = next((p for p in reversed(document.paragraphs) if "公司" in p.text), None)
    if sign_company:
        _set_paragraph_text(sign_company, agent_company)
    sign_date = next(
        (
            p
            for p in reversed(document.paragraphs)
            if "年" in p.text and "月" in p.text and "日" in p.text
        ),
        None,
    )
    if sign_date:
        now = datetime.now()
        _set_paragraph_text(sign_date, f"{now.year}年{now.month}月{now.day}日")
    sign_contact = next(
        (p for p in reversed(document.paragraphs) if any(ch.isdigit() for ch in p.text) and "年" not in p.text),
        None,
    )
    if sign_contact:
        contact_text = f"联系人:{agent_name} 电话:{agent_contact}" if agent_name or agent_contact else ""
        _set_paragraph_text(sign_contact, contact_text)

    for table, crew in (
        (document.tables[0], up),
        (document.tables[1], down),
    ):
        _resize_crew_table(table, len(crew), header_rows=1)
        for index, person in enumerate(crew, start=1):
            values = [
                index,
                person.name,
                person.nationality,
                person.gender,
                _date_text(person.birth_date, "-"),
                person.document_no,
                getattr(person, "luggage", "") or "",
                getattr(person, "temperature", "") or "",
            ]
            for cell, value in zip(table.rows[index].cells, values):
                cell.text = _text(value)
        _apply_table_borders(table)

    output = EXPORT_DIR / f"crew_change_customs_{voyage.id}_{datetime.now():%Y%m%d%H%M%S}.docx"
    document.save(output)
    return output


def _customs_datetime(value, include_year=True):
    if not value:
        return ""
    date_part = value.strftime("%Y年%m月%d日") if include_year else f"{value.month}月{value.day}日"
    return f"{date_part}{value:%H%M}"
